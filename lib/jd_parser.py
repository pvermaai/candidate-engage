"""
JD document parser — extracts structured job description data from uploaded files.
Supports PDF, Word (.docx), and plain text (.txt) formats.
Pipeline: file → raw text → sanitize → Claude structured extraction → validate.
"""

import json
import logging
import re
import pdfplumber
import anthropic
from docx import Document as DocxDocument
from lib.prompts import JD_EXTRACTION_PROMPT
from lib.database import log_api_usage

logger = logging.getLogger(__name__)

MAX_JD_CHARS = 8000
MAX_RETRIES = 2

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def allowed_jd_file(filename: str) -> bool:
    if not filename:
        return False
    ext = _get_extension(filename)
    return ext in ALLOWED_EXTENSIONS


def _get_extension(filename: str) -> str:
    from os.path import splitext
    return splitext(filename.lower())[1]


def _sanitize_text(text: str) -> str:
    """Normalize Unicode chars that cause issues in downstream processing."""
    text = text.replace("\u2013", "-").replace("\u2014", "-")  # en/em dash → hyphen
    text = text.replace("\u2018", "'").replace("\u2019", "'")  # smart quotes
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2192", "->")  # → arrow
    text = text.replace("\u2022", "-")   # bullet
    text = text.replace("\ufffd", "")    # replacement char
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)  # control chars
    return text


def extract_text_from_file(file_path: str, filename: str) -> str:
    ext = _get_extension(filename)
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)
    raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _strip_json_fences(text: str) -> str:
    """Remove markdown code fences wrapping JSON."""
    text = text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()


def _validate_and_normalize(parsed: dict) -> dict:
    """Ensure all expected fields have correct types."""
    for arr_field in ("must_have", "good_to_have", "soft_skills", "responsibilities"):
        if not isinstance(parsed.get(arr_field), list):
            parsed[arr_field] = []

    if parsed.get("experience_min") is not None:
        try:
            parsed["experience_min"] = int(parsed["experience_min"])
        except (ValueError, TypeError):
            parsed["experience_min"] = None

    if parsed.get("experience_max") is not None:
        try:
            parsed["experience_max"] = int(parsed["experience_max"])
        except (ValueError, TypeError):
            parsed["experience_max"] = None

    for str_field in ("title", "location", "mode", "experience", "department"):
        val = parsed.get(str_field)
        if val is not None and not isinstance(val, str):
            parsed[str_field] = str(val)

    return parsed


def extract_jd_fields(raw_text: str) -> dict:
    """Use Claude to extract structured JD fields from raw text, with retry."""
    client = anthropic.Anthropic()

    truncated = raw_text
    if len(raw_text) > MAX_JD_CHARS:
        truncated = raw_text[:MAX_JD_CHARS]
        logger.warning(
            "JD text truncated from %d to %d chars", len(raw_text), MAX_JD_CHARS
        )

    last_error = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                temperature=0,
                messages=[{"role": "user", "content": JD_EXTRACTION_PROMPT + truncated}],
            )

            log_api_usage(
                "jd_extraction",
                response.usage.input_tokens,
                response.usage.output_tokens,
                "claude-sonnet-4-20250514",
            )

            raw = _strip_json_fences(response.content[0].text)
            parsed = json.loads(raw)
            return _validate_and_normalize(parsed)

        except json.JSONDecodeError:
            last_error = "AI returned an invalid response"
            logger.warning("JD extraction attempt %d/%d: bad JSON — %s",
                           attempt, MAX_RETRIES, raw[:200])
        except anthropic.APIStatusError as e:
            last_error = f"AI service error: {e.message}"
            logger.warning("JD extraction attempt %d/%d: API error %s",
                           attempt, MAX_RETRIES, e.status_code)
        except anthropic.APIConnectionError:
            last_error = "Could not reach the AI service. Check your network."
            logger.warning("JD extraction attempt %d/%d: connection error",
                           attempt, MAX_RETRIES)

    raise ValueError(
        f"{last_error}. Tried {MAX_RETRIES} times. Please try again or use manual entry."
    )


def parse_jd_document(file_path: str, filename: str) -> tuple[str, dict]:
    """Full pipeline: file → text → sanitize → structured JD fields."""
    logger.info("Parsing JD document: %s", filename)
    raw_text = extract_text_from_file(file_path, filename)
    if not raw_text.strip():
        raise ValueError("Could not extract any text from the uploaded file")
    sanitized = _sanitize_text(raw_text)
    fields = extract_jd_fields(sanitized)
    return sanitized, fields
